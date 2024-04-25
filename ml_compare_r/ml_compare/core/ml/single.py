# Импорт библиотек и их модулей
import re
import docx
import gensim.downloader as api
import numpy as np
import pandas as pd
from gensim.models import KeyedVectors
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics import pairwise
#from transformers import BertTokenizer, BertModel
from sklearn.metrics.pairwise import cosine_similarity


def parse_paragraphs_to_dataframe(doc):
    all_paragraphs = []
    # Начальный заголовок для текста без явного разделения
    current_header = "Без заголовка"

    for para in doc.paragraphs:
        if para.text.strip():  # Проверка на непустой текст
            # Ищем пункты или заголовки, определяем по форматированию или содержимому
            if para.style.name.startswith('Heading') or (para.text.split()[0].replace('.', '').isdigit() and len(para.text.split()) > 1):
                current_header = para.text.split(maxsplit=1)[0] if para.text.split()[
                    0].replace('.', '').isdigit() else para.text
            all_paragraphs.append((current_header, para.text))

    return pd.DataFrame(all_paragraphs, columns=['Заголовок', 'Текст параграфа'])


def add_section_numbers(df, text_column):
    # Компилируем регулярное выражение для поиска номеров пунктов
    pattern = re.compile(r'^(\d+(\.\d+)*)\s')

    # Функция для извлечения номера пункта из строки
    def extract_number(text):
        match = pattern.match(text)
        return match.group(1) if match else None

    # Применяем функцию к каждому элементу в столбце текста и создаем новую колонку
    df['Номер пункта'] = df[text_column].apply(extract_number)
    return df

# Парсим абзацы в DataFrame


# Обработка текста
def preprocess_text(df, text_column):
    df[text_column] = df[text_column].apply(lambda x: re.sub(
        r'[^\w\s]', '', x.lower()).replace('\n', ' ').strip())
    df[text_column] = df[text_column].apply(lambda x: re.sub(r'\s+', ' ', x))
    return df


def remove_short_requirements(df, column_name='текст требования'):
    """
    Удаляет строки из DataFrame, если количество слов в указанной колонке меньше 5.

    Параметры:
    df (pd.DataFrame): исходный DataFrame, из которого требуется удалить строки.
    column_name (str): имя колонки, в которой проверяется количество слов.

    Возвращает:
    pd.DataFrame: DataFrame после удаления строк.
    """
    # Фильтрация DataFrame: оставляем только те строки, где количество слов в column_name >= 5
    filtered_df = df[df[column_name].apply(lambda x: len(str(x).split()) >= 5)]

    return filtered_df


# Функция для категоризации сходства
def categorize_similarity(similarity):
    if similarity >= 0.1:
        return "Соответствует"
    elif similarity >= 0.07:
        return "Частично соответствует"
    else:
        return "Не соответствует"


def process_files(doc_file, doc_req_file, file_path):
    doc = docx.Document(doc_file)
    doc_req = docx.Document(doc_req_file)
    df_all_text = parse_paragraphs_to_dataframe(doc)
    df_all_text_req = parse_paragraphs_to_dataframe(doc_req)

    # Создаем колонку с номерами пунктов
    df_all_text_req = add_section_numbers(df_all_text_req, 'Текст параграфа')
    df_all_text = add_section_numbers(df_all_text, 'Текст параграфа')
    # Загрузка документа Word

    # Предварительная обработка текста в DataFrame
    df_all_text = preprocess_text(df_all_text, 'Текст параграфа')
    df_all_text_req = preprocess_text(df_all_text_req, 'Текст параграфа')

    # Удаляем строки содержащие менее 5 слов
    df_all_text_req = remove_short_requirements(
        df_all_text_req, 'Текст параграфа')
    df_all_text = remove_short_requirements(df_all_text, 'Текст параграфа')
    # Обновляем индекссацию после удаления
    df_all_text_req = df_all_text_req.reset_index(drop=True)
    df_all_text = df_all_text.reset_index(drop=True)

    # Подготовка текста: соберем все тексты параграфов в список
    texts = df_all_text['Текст параграфа'].tolist()

    # Подготовка списка текстов требований из предыдущей предобработки (словаря требований)
    requirement_texts = df_all_text_req['Текст параграфа'].tolist()

    # Объединение текстов требований и текстов параграфов для обучения одного TF-IDF векторизатора
    combined_texts = texts + requirement_texts

    # Создание и обучение TF-IDF векторизатора на объединенном корпусе
    combined_tfidf_vectorizer = TfidfVectorizer()
    combined_tfidf_matrix = combined_tfidf_vectorizer.fit_transform(
        combined_texts)

    # Разделение матрицы TF-IDF на часть параграфов и часть требований
    tfidf_paragraphs = combined_tfidf_matrix[:len(texts), :]
    tfidf_requirements = combined_tfidf_matrix[len(texts):, :]

    # Расчет косинусного сходства между параграфами и требованиями
    cosine_similarities = cosine_similarity(
        tfidf_paragraphs, tfidf_requirements)

    # Инициализация пустого DataFrame для наилучших совпадений
    best_match_df = pd.DataFrame(columns=['Пункт требования', 'Текст требования',
                                          'Заголовок параграфа', 'Пункт параграфа', 'Текст параграфа', 'Сходство'])

    rows_to_add = []
    for req_index, req_row in df_all_text_req.iterrows():
        paragraph_index = np.argmax(cosine_similarities[:, req_index])
        max_similarity = cosine_similarities[paragraph_index, req_index]
        paragraph_data = df_all_text.loc[paragraph_index]
        paragraph_data_req = df_all_text_req.loc[paragraph_index]
        # Создание временного DataFrame для текущей строки
        temp_df = pd.DataFrame({
            'Заголовок требования': [req_row['Заголовок']],
            'Номер пункта требования': [req_row['Номер пункта']],
            'Текст требования': [req_row['Текст параграфа']],
            'Заголовок параграфа': [paragraph_data['Заголовок']],
            # Предполагая, что 'Номер пункта' есть в df_all_text
            'Номер пункта параграфа': [paragraph_data['Номер пункта']],
            'Текст параграфа': [paragraph_data['Текст параграфа']],
            'Сходство': [max_similarity]
        })
        rows_to_add.append(temp_df)

    # Добавление всех собранных строк в основной DataFrame
    best_match_df = pd.concat(rows_to_add, ignore_index=True)
    # Применение функции к колонке 'Сходство'
    best_match_df['Степень соответствия'] = best_match_df['Сходство'].apply(
        categorize_similarity)
    best_match_df.to_excel(file_path, index=False, engine='openpyxl')
